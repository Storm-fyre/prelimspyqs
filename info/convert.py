#!/usr/bin/env python3
"""
DOCX/TXT → JSON converter for UPSC-style Q&A with year headings.

- Reads .docx and .txt.
- Preserves line breaks in question stems.
- Converts every Word TABLE into one Markdown table block (with leading/trailing pipes),
  then embeds it inside the question stem (so rows don't split into separate questions).
- Recognizes options starting with (a), a., or a) — LOWERCASE a–d only (as requested).
- Safely handles duplicate option labels (e.g., two '(a)') by putting later one
  in the next free slot (b/c/d) instead of overwriting.
- Starts a new question on "N." only if we haven't started one yet OR we've already
  seen any options for the current question (prevents numbered sub-points from splitting).
- Asks for input/output filenames; coerces output ".docx" to ".json" to avoid overwrite.
- Outputs difficulty "M", empty answer/explanation.
"""

import re
import json
import sys
from pathlib import Path

# ---------- Regex ----------
YEAR_RE = re.compile(r'^\s*(\d{4})\s*$')
QSTART_RE = re.compile(r'^\s*(\d+)\.\s*(.*\S)?\s*$')  # e.g., "3. Consider..."
OPT_RE = re.compile(r'^\s*\(?([a-d])[\).]\s*(.*\S)?\s*$')  # (a) / a. / a)  (lowercase only)

# ---------- Text utils ----------
def smart_quotes_to_ascii(s: str) -> str:
    repl = {
        "\u2018": "'", "\u2019": "'",
        "\u201C": '"', "\u201D": '"',
        "\u2013": "-", "\u2014": "-",
        "\xa0": " ",
    }
    for k, v in repl.items():
        s = s.replace(k, v)
    return s

def finalize_question_text(lines):
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    out, prev_blank = [], False
    for ln in lines:
        if ln.strip() == '':
            if not prev_blank:
                out.append('')
            prev_blank = True
        else:
            out.append(ln.rstrip())
            prev_blank = False
    return "\n".join(out)

def normalize_inline(s: str) -> str:
    return re.sub(r'\s+', ' ', s).strip()

# ---------- DOCX reading with tables in original order ----------
def iter_block_items(doc):
    from docx.document import Document as _Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    def _iter_block_items(parent):
        for child in parent.element.body.iterchildren():
            if child.tag.endswith('p'):
                yield Paragraph(child, parent)
            elif child.tag.endswith('tbl'):
                yield Table(child, parent)

    if isinstance(doc, _Document):
        for block in _iter_block_items(doc):
            yield block
    else:
        for p in doc.paragraphs:
            yield p

def table_to_markdown(tbl) -> str:
    rows = []
    for row in tbl.rows:
        cells = []
        for cell in row.cells:
            txt = " ".join(smart_quotes_to_ascii(p.text.strip()) for p in cell.paragraphs if p.text is not None)
            txt = re.sub(r'\s+', ' ', txt).strip()
            cells.append(txt)
        rows.append(cells)
    if not rows:
        return ""
    col_count = max(len(r) for r in rows)
    rows = [r + [""] * (col_count - len(r)) for r in rows]
    header = "| " + " | ".join(rows[0]) + " |"
    sep = "| " + " | ".join("---" for _ in range(col_count)) + " |"
    lines = [header, sep]
    for r in rows[1:]:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)

def read_docx_text(path: Path) -> str:
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError:
        print("The 'python-docx' package is required. Install with: pip install python-docx")
        sys.exit(1)

    doc = Document(str(path))
    blocks = []
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = smart_quotes_to_ascii((block.text or "").replace("\r", "\n"))
            blocks.append(text)
        elif isinstance(block, Table):
            md = table_to_markdown(block)
            if md.strip():
                blocks.append(md)
    text = "\n".join(blocks)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text

def read_input_text(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        return read_docx_text(path)
    else:
        return smart_quotes_to_ascii(path.read_text(encoding="utf-8", errors="ignore"))

# ---------- Core parser ----------
def parse_text(text: str):
    years = {}
    current_year = None
    current_qnum = None
    q_lines = []
    options = {}
    current_opt = None
    in_options = False

    def place_option(letter: str, text_val: str):
        if letter in options and options[letter].strip():
            for fallback in ['a', 'b', 'c', 'd']:
                if fallback not in options or not options[fallback].strip():
                    letter = fallback
                    break
        options[letter] = text_val

    def flush_question():
        nonlocal current_qnum, q_lines, options, current_opt, in_options, current_year
        if current_qnum is None or current_year is None:
            return
        question_text = finalize_question_text(q_lines)
        opt_list = []
        for letter in ['a', 'b', 'c', 'd']:
            t = normalize_inline(options.get(letter, ''))
            opt_list.append(f"{letter.upper()}. {t}" if t else f"{letter.upper()}. ")
        q_id = f"POL-{current_year}-{int(current_qnum):02d}"
        years.setdefault(str(current_year), []).append({
            "id": q_id,
            "question": question_text,
            "options": opt_list,
            "answer": "",
            "explanation": "",
            "difficulty": "M"
        })
        current_qnum = None
        q_lines = []
        options = {}
        current_opt = None
        in_options = False

    for raw in text.splitlines():
        line = raw.rstrip()

        # Year heading
        m_year = YEAR_RE.match(line)
        if m_year:
            flush_question()
            current_year = m_year.group(1)
            continue

        # Question start "N.":
        # Start a new question if (a) we aren't in one yet, OR (b) we've already seen options
        # for the current one (so numbered sub-points earlier don't split questions).
        m_q = QSTART_RE.match(line) if current_year else None
        if m_q:
            if current_qnum is None or options:  # <-- key change
                flush_question()
                current_qnum = m_q.group(1)
                first_part = (m_q.group(2) or '').strip()
                if first_part:
                    q_lines.append(first_part)
                in_options = False
                current_opt = None
                continue
            else:
                # still inside stem before options; treat as numbered sub-point
                q_lines.append(line)
                continue

        # Options (a)/(a.)/(a)) lowercase only
        if current_year and current_qnum is not None:
            m_opt = OPT_RE.match(line)
            if m_opt:
                in_options = True
                current_opt = m_opt.group(1)
                opt_text = (m_opt.group(2) or '').strip()
                place_option(current_opt, opt_text)
                continue

            # Continuations
            if in_options and current_opt:
                if line.strip():
                    options[current_opt] = (options.get(current_opt, '') + ' ' + line.strip()).strip()
            else:
                q_lines.append(line)

    flush_question()
    return years

# ---------- Main ----------
def main():
    in_path = input("Enter input filename (.docx or .txt): ").strip()
    out_path = input("Enter output JSON filename: ").strip()
    out_p = Path(out_path)
    if out_p.suffix.lower() == ".docx":
        out_p = out_p.with_suffix(".json")

    in_file = Path(in_path)
    if not in_file.exists():
        print(f"Error: input file '{in_file}' not found.")
        return

    text = read_input_text(in_file)
    years = parse_text(text)

    data = {"subject": "Polity", "years": years}
    out_p.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding='utf-8')
    print(f"Done. Wrote {out_p}")

if __name__ == "__main__":
    main()
